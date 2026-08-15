"""Word 리포트 생성 테스트 — 화면과 같은 숫자가 문서에 들어가는지 확인한다."""

from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document

from src import report, valuation
from src.models import CanslimItem, CanslimResult, Verdict
from src.tickers import StockRef


@pytest.fixture
def analysis(snapshot, quarterly, annual):
    """네트워크 없이 삼성전자 실측 fixture 로 Analysis 모양을 만든다."""
    val = valuation.compute_valuation(snapshot, quarterly, annual)
    canslim = CanslimResult(items=[
        CanslimItem("C", "최근 분기 실적", "분기 EPS 전년동기比 +25% 이상",
                    "+489.6%", Verdict.PASS, "2026.03 6,993원 vs 2025.03 1,186원"),
        CanslimItem("A", "연간 실적", "연간 EPS 3년 CAGR +25% 이상",
                    "+75.5%", Verdict.PASS, ""),
        CanslimItem("I", "수급", "기관/외국인 보유 증가", "—", Verdict.UNKNOWN, "데이터 부족"),
    ])
    return SimpleNamespace(
        ref=StockRef(code="005930", name="삼성전자", market="KOSPI", region="KR"),
        snap=snapshot, quarterly=quarterly, annual=annual,
        valuation=val, canslim=canslim,
    )


def all_text(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class Test리포트:
    def test_docx로_열리고_제목이_있다(self, analysis):
        data = report.build_report(analysis)
        doc = Document(BytesIO(data))
        assert "삼성전자 분석 리포트" in doc.paragraphs[0].text

    def test_화면과_같은_핵심_숫자가_들어간다(self, analysis):
        text = all_text(report.build_report(analysis))
        assert "254,000원" in text          # 현재가
        assert "+489.6%" in text            # CANSLIM C 실제값
        assert "PER" in text and "PEG" in text
        assert "2026.03" in text            # 분기 실적 표
        assert "6,993원" in text            # 분기 EPS

    def test_판정과_출처_구분이_글자로_들어간다(self, analysis):
        text = all_text(report.build_report(analysis))
        assert "합격" in text and "판단불가" in text
        assert "실적확정" in text and "컨센서스" in text

    def test_역산_분기가_있으면_역산추정_행이_있다(self, analysis):
        assert analysis.valuation.derived is not None  # 삼성 fixture 는 역산이 성립한다
        text = all_text(report.build_report(analysis))
        assert "역산추정" in text

    def test_면책_문구가_있다(self, analysis):
        text = all_text(report.build_report(analysis))
        assert "매수·매도 신호가 아닙니다" in text

    def test_한글_글꼴을_동아시아_폰트로_지정한다(self, analysis):
        """w:eastAsia 를 지정하지 않으면 한글이 기본 세리프 글꼴로 렌더링된다."""
        doc = Document(BytesIO(report.build_report(analysis)))
        rpr = doc.styles["Normal"].element.rPr
        assert rpr.rFonts.get(report.qn("w:eastAsia")) == report.MALGUN
