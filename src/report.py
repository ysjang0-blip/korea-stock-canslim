"""분석 결과를 Word(.docx) 리포트로 만든다.

차트는 넣지 않는다 — 글과 표만. 화면(app.py)이 보여주는 것과 같은 숫자를
같은 서식으로 담아, 화면과 리포트가 다른 말을 하지 않게 한다.
"""

from __future__ import annotations

import datetime as dt
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from .models import Source, Verdict

MALGUN = "맑은 고딕"
MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_SOURCE_MARK = {
    Source.ACTUAL: "실적확정",
    Source.CONSENSUS: "컨센서스",
    Source.DERIVED: "역산추정",
}


def _korean_style(style, size: float | None = None) -> None:
    """스타일 글꼴을 맑은 고딕으로. 라틴(w:ascii)만 바꾸면 한글은 기본 글꼴로 남는다."""
    style.font.name = MALGUN
    rpr = style.element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), MALGUN)
    if size is not None:
        style.font.size = Pt(size)


def _new_document() -> Document:
    doc = Document()
    _korean_style(doc.styles["Normal"], 10)
    for name in ("Title", "Heading 1", "Heading 2", "List Bullet"):
        if name in doc.styles:
            _korean_style(doc.styles[name])
    return doc


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        for cell, text in zip(table.add_row().cells, row):
            cell.text = text
    return table


def build_report(a) -> bytes:
    """Analysis 객체 하나를 받아 완성된 .docx 바이트를 돌려준다."""
    snap, val = a.snap, a.valuation
    usd = snap.currency == "USD"
    doc = _new_document()

    # ── 표지 ──────────────────────────────────────────────────────────
    doc.add_heading(f"{snap.name} 분석 리포트", level=0)
    doc.add_paragraph(f"{a.ref.code} · {a.ref.market} · CANSLIM + 밸류에이션")

    latest_actual = a.quarterly.actual_periods()
    bits = [f"가격 기준일 {snap.price_date_label or '—'}"]
    if latest_actual:
        bits.append(f"재무 {latest_actual[-1].label} 분기까지 확정")
    if snap.consensus_date:
        bits.append(f"컨센서스 {snap.consensus_date}")
    bits.append(f"리포트 생성 {dt.date.today().isoformat()}")
    doc.add_paragraph(" · ".join(bits))

    # ── 현재 시세 ─────────────────────────────────────────────────────
    doc.add_heading("현재 시세", level=1)
    change_txt = ""
    if snap.change is not None and snap.change_pct is not None:
        change_txt = (f" (전일 대비 {snap.change:+,.2f}달러, {snap.change_pct:+.2f}%)" if usd
                      else f" (전일 대비 {snap.change:+,.0f}원, {snap.change_pct:+.2f}%)")
    upside = (f" (상승여력 {snap.upside_pct:+.1f}%)" if snap.upside_pct is not None else "")
    ratio_52w = (f"{snap.price / snap.high_52w:.0%}"
                 if snap.price and snap.high_52w else "—")
    _add_table(doc, ["항목", "값"], [
        ["현재가", (snap.money(snap.price) if snap.price else "—") + change_txt],
        ["시가총액", snap.money_big(snap.market_cap) if snap.market_cap else "—"],
        ["목표주가 평균", (snap.money(snap.target_price) if snap.target_price else "—") + upside],
        ["투자의견", f"{snap.recomm_mean:.2f} / 5 (5에 가까울수록 매수 우세)"
         if snap.recomm_mean else "—"],
        ["52주 최고 / 최저", f"{snap.money(snap.high_52w)} / {snap.money(snap.low_52w)}"
         if snap.high_52w and snap.low_52w else "—"],
        ["52주 최고 대비", ratio_52w],
    ])

    # ── CANSLIM ──────────────────────────────────────────────────────
    doc.add_heading("CANSLIM 판정", level=1)
    doc.add_paragraph(f"{a.canslim.summary} — 판단이 불가능한 항목은 분모에서 제외했습니다.")
    _add_table(
        doc,
        ["항목", "기준", "실제값", "판정", "근거"],
        [[f"{i.letter} {i.name}", i.criterion, i.actual, i.verdict.value, i.evidence]
         for i in a.canslim.items],
    )

    # ── 밸류에이션 ────────────────────────────────────────────────────
    doc.add_heading("밸류에이션", level=1)
    if usd:
        eps_unit, eps_digits = "$", 2
        rev_unit, rev_digits, rev_scale = "억$", 0, 1e8
    else:
        eps_unit, eps_digits = "원", 0
        rev_unit, rev_digits, rev_scale = "조원", 1, 1e12

    def metric_cells(attr: str, unit: str, digits: int = 2, scale: float = 1.0) -> list[str]:
        cells = []
        for c in val.columns:
            m = getattr(c, attr)
            cells.append("—" if not m.is_ok else f"{m.value / scale:,.{digits}f}{unit}")
        return cells

    headers = ["지표"] + [f"{c.label} ({_SOURCE_MARK.get(c.source, c.source.value)})"
                          for c in val.columns]
    _add_table(doc, headers, [
        ["PER"] + metric_cells("per", "배"),
        ["PSR"] + metric_cells("psr", "배"),
        ["PEG"] + metric_cells("peg", ""),
        ["ROE"] + metric_cells("roe", "%"),
        ["EPS (12개월)"] + metric_cells("eps_ttm", eps_unit, eps_digits),
        ["매출 (12개월)"] + metric_cells("revenue_ttm", rev_unit, rev_digits, rev_scale),
        ["산출 근거"] + [c.note for c in val.columns],
    ])

    g = val.growth
    notes = doc.add_paragraph()
    notes.add_run(
        f"교차검증 — {val.per_cross_check}\n"
        f"PEG에 쓴 성장률 — 현재 열은 연간 EPS {g.annual_cagr_years}년 CAGR "
        f"{g.annual_cagr.text('%')}, 예상 열은 연간 컨센서스 성장률 {g.forward_annual.text('%')}.\n"
        f"최근 분기 EPS — 전년 동기 대비 {g.quarter_yoy.text('%')}"
    ).font.size = Pt(9)

    # ── 분기 실적 ─────────────────────────────────────────────────────
    doc.add_heading("분기 실적", level=1)
    rev_q_unit, rev_q_digits, rev_q_scale = (("억$", 1, 1e8) if usd else ("억원", 0, 1.0))

    def fmt(value: float | None, unit: str, digits: int, scale: float = 1.0) -> str:
        return "—" if value is None else f"{value / scale:,.{digits}f}{unit}"

    q_rows = []
    for p in a.quarterly.periods:
        eps, rev = a.quarterly.value("EPS", p.key), a.quarterly.value("매출액", p.key)
        if eps is None and rev is None:
            continue
        source = Source.CONSENSUS if p.is_consensus else Source.ACTUAL
        q_rows.append([p.label, fmt(eps, eps_unit, eps_digits),
                       fmt(rev, rev_q_unit, rev_q_digits, rev_q_scale), _SOURCE_MARK[source]])
    if val.derived:
        d = val.derived
        q_rows.append([f"{d.key[:4]}.{d.key[4:]}",
                       fmt(d.values.get("EPS"), eps_unit, eps_digits),
                       fmt(d.values.get("매출액"), rev_q_unit, rev_q_digits, rev_q_scale),
                       f"{_SOURCE_MARK[Source.DERIVED]} ({d.method})"])
    _add_table(doc, ["분기", "EPS", "매출액", "구분"], q_rows)

    # ── 한계와 면책 ───────────────────────────────────────────────────
    doc.add_heading("이 분석의 한계", level=1)
    limits = [
        "오닐 원본 RS Rating(1~99점)이 아니라 시장 지수 대비 초과수익 여부로 L을 판정합니다.",
        "N(New)의 재료는 뉴스 제목을 키워드로 자동 분류한 것이라 사람의 판단을 대신할 수 없습니다.",
        "S는 유통주식수(float)를 반영하지 못해 거래량 급증으로 대체했습니다.",
    ]
    if usd:
        limits += [
            "I는 기관 보유 비중의 현재 값만 있어 '증가 추세'는 판정할 수 없습니다.",
            "야후 파이낸스 비공식 라이브러리(yfinance)에 의존해, 형식이 바뀌면 값이 달라질 수 있습니다.",
        ]
    else:
        limits += [
            "Q+2는 연간 컨센서스에서 역산한 추정치라 신뢰도가 한 단계 낮습니다.",
            "네이버 비공식 API에 의존해, 형식이 바뀌면 값이 달라질 수 있습니다.",
        ]
    for line in limits:
        doc.add_paragraph(line, style="List Bullet")

    disclaimer = doc.add_paragraph()
    run = disclaimer.add_run(
        f"데이터: {snap.source_name} · 이 리포트는 투자 판단을 돕는 참고 자료이며 "
        "매수·매도 신호가 아닙니다. 투자 결정과 그 결과는 전적으로 본인의 책임입니다."
    )
    run.bold = True
    run.font.size = Pt(9)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
